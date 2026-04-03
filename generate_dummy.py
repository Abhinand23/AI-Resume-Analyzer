import docx
doc = docx.Document()
doc.add_heading('John Doe - Senior Software Engineer', 0)
doc.add_paragraph('Skills: Python, React, FastAPI, SQL, Node.js, Frontend Design')
doc.add_paragraph('Experience: 5 years building scalable web applications, REST APIs, and modern frontends for luxury brands.')
doc.save('dummy_resume.docx')
